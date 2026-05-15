from __future__ import annotations

from io import BytesIO

from docx import Document


def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """
    Extract plain text from a DOCX file (bytes).
    """
    bio = BytesIO(docx_bytes)
    doc = Document(bio)

    parts: list[str] = []

    # Paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # Tables (optional but useful)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                ct = (cell.text or "").strip()
                if ct:
                    row_text.append(ct)
            if row_text:
                parts.append(" | ".join(row_text))

    return "\n".join(parts).strip()
