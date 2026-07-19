from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from io import BytesIO
from docx import Document
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter

import json

from sqlalchemy.orm import Session

from fastapi import (
    Depends,
    HTTPException
)

from api.db.database import get_db

from api.models.sermon import Sermon

from api.core.dependencies import get_current_user

from api.models.user import User
from api.models.shared_sermon import SharedSermon

# =========================================
# SCHEMAS
# =========================================
from api.schemas.faith import (
    SermonRequest,
    SermonResponse,
    SermonRefineRequest
)

# =========================================
# AI SERVICES
# =========================================
from api.services.ai_service import (
    generate_ai_response,
    refine_sermon,
    generate_multiverse_sermons
)
from api.models.sermon_comment import SermonComment

router = APIRouter()


# =========================================
# GENERATE SERMON
# =========================================
@router.post(
    "/sermon",
    response_model=SermonResponse
)
async def create_sermon(
    payload: SermonRequest,
    current_user: User = Depends(
        get_current_user
    )
):
    # =====================================
    # MEMBER + PASTOR ACCESS
    # =====================================
    if current_user.role not in [
        "member",
        "pastor",
        "admin"
    ]:
        raise HTTPException(
            status_code=403,
            detail="Unauthorized"
        )
    # =====================================

    print("\n==========")
    print("SERMON REQUEST")
    print("Scripture:", payload.scripture)
    print("Input:", payload.input)
    print("==========\n")
    sermon = generate_ai_response(
    payload
    )

    # =====================================
    # FORCE REQUESTED BIBLE VERSE
    # =====================================

    if payload.scripture:

        sermon["scripture"] = payload.scripture
    # =====================================
    # COMPATIBILITY LAYER
    # =====================================

    if (
        "points" in sermon
        and "main_points" not in sermon
    ):

        sermon["main_points"] = sermon["points"]

    print(
        "REQUESTED BIBLE:",
        payload.scripture
    )

    print(
        "RETURNED SCRIPTURE:",
        sermon.get("scripture")
    )

    return sermon

# =========================================
# REFINE SERMON
# =========================================
@router.post("/sermon/refine")
async def refine_existing_sermon(

    payload: SermonRefineRequest,

    current_user: User = Depends(
        get_current_user
    )
):

    try:

        # =========================
        # PASTOR ONLY
        # =========================
        if current_user.role not in [
            "pastor",
            "admin"
        ]:

            raise HTTPException(
                status_code=403,
                detail="Pastor verification required"
            )

        # =========================
        # VALIDATE INPUT
        # =========================
        if not payload.sermon:

            raise HTTPException(
                status_code=400,
                detail="Sermon content required"
            )

        if not payload.refine_type:

            raise HTTPException(
                status_code=400,
                detail="Refine type required"
            )

        # =========================
        # REFINE
        # =========================
        refined = refine_sermon(
            payload.sermon,
            payload.refine_type
        )

        return {

            "success": True,

            "refine_type":
                payload.refine_type,

            "sermon":
                refined
        }

    except HTTPException:

        raise

    except Exception as e:

        print(
            "REFINE SERMON ERROR:",
            str(e)
        )

        raise HTTPException(
            status_code=500,
            detail=
                "Failed to refine sermon"
        )


# =========================================
# MULTIVERSE SERMONS
# =========================================
@router.post("/sermon/multiverse")
async def multiverse_sermons(
    payload: SermonRequest
):

    universes = generate_multiverse_sermons(
        payload.topic
    )

    return universes


# =========================================
# HEALTH CHECK
# =========================================
@router.get("/sermon-health")
async def sermon_health():

    return {
        "status": "ok",
        "service": "XynaFaith Sermon AI"
    }


# =========================================
# EXPORT PDF
# =========================================
@router.post("/sermon/export-pdf")
async def export_sermon_pdf(
    payload: dict,
    current_user: User = Depends(
        get_current_user
    )
):
    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=50,
        leftMargin=50,
        topMargin=60,
        bottomMargin=60
    )

    styles = getSampleStyleSheet()

    # =====================================
    # PREMIUM STYLES
    # =====================================
    title_style = styles["Title"]
    title_style.fontSize = 28
    title_style.leading = 34

    heading_style = styles["Heading2"]
    heading_style.fontSize = 18
    heading_style.leading = 24

    body_style = styles["BodyText"]
    body_style.fontSize = 12
    body_style.leading = 22

    story = []

    # =========================
    # TITLE
    # =========================
    title = payload.get(
        "title",
        "Untitled Sermon"
    )

    story.append(
        Paragraph(
            f"<b>{title}</b>",
            title_style
        )
    )

    story.append(
        Spacer(1, 24)
    )

    # =========================
    # SCRIPTURE
    # =========================
    scripture = payload.get(
        "scripture"
    )

    if scripture:

        story.append(
            Paragraph(
                f"<i>📖 {scripture}</i>",
                heading_style
            )
        )

        story.append(
            Spacer(1, 24)
        )

    # =========================
    # INTRODUCTION
    # =========================
    intro = payload.get(
        "introduction",
        ""
    )

    story.append(
        Paragraph(
            "<b>Introduction</b>",
            heading_style
        )
    )

    story.append(
        Paragraph(
            intro,
            body_style
        )
    )

    story.append(
        Spacer(1, 24)
    )

    # =========================
    # MAIN POINTS
    # =========================
    points = payload.get(
        "main_points",
        []
    )

    for point in points:

        story.append(
            Paragraph(
                f"<b>{point.get('title', '')}</b>",
                heading_style
            )
        )

        story.append(
            Paragraph(
                point.get(
                    "content",
                    ""
                ),
                body_style
            )
        )

        story.append(
            Spacer(1, 24)
        )

    # =========================
    # APPLICATION
    # =========================
    application = payload.get(
        "application",
        ""
    )

    story.append(
        Paragraph(
            "<b>Application</b>",
            heading_style
        )
    )

    story.append(
        Paragraph(
            application,
            body_style
        )
    )

    story.append(
        Spacer(1, 24)
    )

    # =========================
    # CONCLUSION
    # =========================
    conclusion = payload.get(
        "conclusion",
        ""
    )

    story.append(
        Paragraph(
            "<b>Conclusion</b>",
            heading_style
        )
    )

    story.append(
        Paragraph(
            conclusion,
            body_style
        )
    )

    # =========================
    # FOOTER BRANDING
    # =========================
    story.append(
        Spacer(1, 40)
    )

    story.append(
        Paragraph(
            "Generated with XynaFaith",
            styles["Italic"]
        )
    )

    # =========================
    # BUILD PDF
    # =========================
    doc.build(story)

    buffer.seek(0)

    filename = (
        title
        .replace(" ", "_")
        + ".pdf"
    )

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={
            "Content-Disposition":
            f"attachment; filename={filename}"
        }
    )


# =========================================
# EXPORT DOCX
# =========================================
@router.post("/sermon/export-docx")
async def export_sermon_docx(
    payload: dict,
    current_user: User = Depends(
        get_current_user
    )
):
    document = Document()

    # =========================
    # TITLE
    # =========================
    title = payload.get(
        "title",
        "Untitled Sermon"
    )

    document.add_heading(
        title,
        level=1
    )

    # =========================
    # SCRIPTURE
    # =========================
    scripture = payload.get(
        "scripture"
    )

    if scripture:

        document.add_heading(
            "Scripture",
            level=2
        )

        document.add_paragraph(
            scripture
        )

    # =========================
    # INTRODUCTION
    # =========================
    intro = payload.get(
        "introduction",
        ""
    )

    document.add_heading(
        "Introduction",
        level=2
    )

    document.add_paragraph(
        intro
    )

    # =========================
    # MAIN POINTS
    # =========================
    points = payload.get(
        "main_points",
        []
    )

    for point in points:

        document.add_heading(
            point.get(
                "title",
                "Point"
            ),
            level=2
        )

        document.add_paragraph(
            point.get(
                "content",
                ""
            )
        )

    # =========================
    # APPLICATION
    # =========================
    application = payload.get(
        "application",
        ""
    )

    document.add_heading(
        "Application",
        level=2
    )

    document.add_paragraph(
        application
    )

    # =========================
    # CONCLUSION
    # =========================
    conclusion = payload.get(
        "conclusion",
        ""
    )

    document.add_heading(
        "Conclusion",
        level=2
    )

    document.add_paragraph(
        conclusion
    )

    # =========================
    # FOOTER
    # =========================
    document.add_paragraph(
        "\nGenerated with XynaFaith"
    )

    # =========================
    # SAVE TO MEMORY
    # =========================
    buffer = BytesIO()

    document.save(
        buffer
    )

    buffer.seek(0)

    filename = (
        title
        .replace(" ", "_")
        + ".docx"
    )

    return StreamingResponse(

        buffer,

        media_type=
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",

        headers={
            "Content-Disposition":
            f"attachment; filename={filename}"
        }
    )

# =========================================
# SAVE SERMON
# =========================================
@router.post("/sermon/save")
async def save_sermon(

    payload: dict,

    db: Session = Depends(
        get_db
    ),

    current_user: User = Depends(
        get_current_user
    )
):

    try:

        # =========================
        # TITLE
        # =========================
        title = payload.get(
            "title",
            "Untitled Sermon"
        )

        # =========================
        # SCRIPTURE
        # =========================
        scripture = payload.get(
            "scripture"
        )

        # =========================
        # CREATE SERMON
        # =========================
        sermon = Sermon(

            title=title,

            scripture=scripture,

            author_id=current_user.id,

            # LEGACY STORAGE
            content=json.dumps(
                payload
            ),

            # NEW STRUCTURED STORAGE
            sermon_data=payload,

            is_public=0
        )

        # =========================
        # SAVE
        # =========================
        db.add(sermon)

        db.commit()

        db.refresh(sermon)

        return {

            "success": True,

            "message":
                "Sermon saved",

            "sermon_id":
                sermon.id
        }

    except Exception as e:

        print(
            "SAVE SERMON ERROR:",
            str(e)
        )

        raise HTTPException(
            status_code=500,
            detail="Failed to save sermon"
        )
    
# =========================================
# GET MY SERMONS
# =========================================
@router.get("/sermon/my")
async def get_my_sermons(

    db: Session = Depends(
        get_db
    ),

    current_user: User = Depends(
        get_current_user
    )
):

    try:

        # =========================
        # FETCH SERMONS
        # =========================
        sermons = (

            db.query(Sermon)

            .filter(
                Sermon.author_id
                == current_user.id
            )

            .order_by(
                Sermon.created_at.desc()
            )

            .all()
        )

        # =========================
        # FORMAT RESPONSE
        # =========================
        formatted = []

        for sermon in sermons:

            formatted.append({

                "id":
                    sermon.id,

                "title":
                    sermon.title,

                "scripture":
                    sermon.scripture,

                "created_at":
                    sermon.created_at,

                "views":
                    sermon.views,

                "shares":
                    sermon.shares,

                "is_public":
                    sermon.is_public
            })

        return {
            "success": True,
            "sermons": formatted
        }

    except Exception as e:

        print(
            "GET MY SERMONS ERROR:",
            str(e)
        )

        raise HTTPException(
            status_code=500,
            detail=
              "Failed to load sermons"
        )
    
# =========================================
# GET SINGLE SERMON
# =========================================
@router.get("/sermon/{sermon_id}")
async def get_single_sermon(
    sermon_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    try:
        # =====================================
        # LOAD SERMON
        # =====================================
        sermon = (
            db.query(Sermon)
            .filter(
                Sermon.id == sermon_id
            )
            .first()
        )

        if not sermon:
            raise HTTPException(
                status_code=404,
                detail="Sermon not found"
            )

        # =====================================
        # ACCESS CONTROL
        # =====================================
        is_owner = (
            sermon.author_id ==
            current_user.id
        )

        is_public = bool(
            sermon.is_public
        )

        is_admin = (
            getattr(
                current_user,
                "role",
                ""
            ) == "admin"
        )

        has_shared_access = False

        if not is_owner:
            shared = (
                db.query(SharedSermon)
                .filter(
                    SharedSermon.sermon_id ==
                    sermon.id,
                    SharedSermon.to_pastor_id ==
                    current_user.id
                )
                .first()
            )

            has_shared_access = (
                shared is not None
            )

        if not (
            is_owner or
            is_public or
            is_admin or
            has_shared_access
        ):
            raise HTTPException(
                status_code=403,
                detail=(
                    "You do not have permission "
                    "to view this sermon"
                )
            )

        # =====================================
        # NORMALIZE SERMON CONTENT
        # =====================================
        sermon_content = {}

        if (
            sermon.sermon_data and
            isinstance(
                sermon.sermon_data,
                dict
            )
        ):
            sermon_content = dict(
                sermon.sermon_data
            )

        elif sermon.content:
            try:
                parsed = json.loads(
                    sermon.content
                )

                if isinstance(
                    parsed,
                    dict
                ):
                    sermon_content = parsed
                else:
                    sermon_content = {
                        "introduction":
                            str(parsed)
                    }

            except (
                json.JSONDecodeError,
                TypeError
            ):
                sermon_content = {
                    "introduction":
                        sermon.content
                }

        # =====================================
        # ENSURE REQUIRED METADATA
        # =====================================
        sermon_content["id"] = (
            sermon.id
        )

        sermon_content["title"] = (
            sermon_content.get(
                "title"
            )
            or sermon.title
            or "Untitled Sermon"
        )

        sermon_content["scripture"] = (
            sermon_content.get(
                "scripture"
            )
            or sermon.scripture
            or ""
        )

        sermon_content["author_id"] = (
            sermon.author_id
        )

        sermon_content["author_name"] = (
            sermon.author.name
            if sermon.author
            else ""
        )

        sermon_content["is_public"] = (
            bool(
                sermon.is_public
            )
        )

        sermon_content["created_at"] = (
            sermon.created_at
        )

        sermon_content["updated_at"] = (
            sermon.updated_at
        )

        sermon_content["views"] = (
            sermon.views or 0
        )

        sermon_content["shares"] = (
            sermon.shares or 0
        )

        # =====================================
        # UPDATE VIEW COUNT
        # =====================================
        sermon.views = (
            sermon.views or 0
        ) + 1

        db.commit()

        # =====================================
        # RESPONSE
        # =====================================
        return {
            "success": True,
            "sermon": sermon_content
        }

    except HTTPException:
        raise

    except Exception as e:
        db.rollback()

        print(
            "GET SERMON ERROR:",
            repr(e)
        )

        raise HTTPException(
            status_code=500,
            detail="Failed to load sermon"
        )
    
# =========================================
# UPDATE SERMON
# =========================================
@router.put("/sermon/update/{sermon_id}")
async def update_sermon(

    sermon_id: int,

    payload: dict,

    db: Session = Depends(
        get_db
    ),

    current_user: User = Depends(
        get_current_user
    )
):

    try:

        sermon = (

            db.query(Sermon)

            .filter(
                Sermon.id == sermon_id
            )

            .filter(
                Sermon.author_id
                == current_user.id
            )

            .first()
        )

        if not sermon:

            raise HTTPException(
                status_code=404,
                detail="Sermon not found"
            )

        # =========================
        # UPDATE FIELDS
        # =========================
        sermon.title = payload.get(
            "title",
            sermon.title
        )

        sermon.scripture = payload.get(
            "scripture",
            sermon.scripture
        )

        # LEGACY
        sermon.content = json.dumps(
            payload
        )

        # STRUCTURED
        sermon.sermon_data = payload

        db.commit()

        db.refresh(sermon)

        return {

            "success": True,

            "message":
                "Sermon updated"
        }

    except Exception as e:

        print(
            "UPDATE SERMON ERROR:",
            str(e)
        )

        raise HTTPException(
            status_code=500,
            detail=
              "Failed to update sermon"
        )
    
# =========================================
# SHARE SERMON
# =========================================
@router.post("/sermon/share")
async def share_sermon(

    payload: dict,

    db: Session = Depends(
        get_db
    ),

    current_user: User = Depends(
        get_current_user
    )
):

    try:

        # =========================
        # PASTOR ONLY
        # =========================
        if current_user.role not in [
            "pastor",
            "admin"
        ]:

            raise HTTPException(
                status_code=403,
                detail="Pastor verification required"
            )

        # =========================
        # INPUTS
        # =========================
        sermon_id = payload.get(
            "sermon_id"
        )

        to_pastor_id = payload.get(
            "to_pastor_id"
        )

        comment = payload.get(
            "comment",
            ""
        )

        # =========================
        # VALIDATE INPUT
        # =========================
        if not sermon_id:

            raise HTTPException(
                status_code=400,
                detail="sermon_id is required"
            )

        if not to_pastor_id:

            raise HTTPException(
                status_code=400,
                detail="to_pastor_id is required"
            )

        # =========================
        # VALIDATE SERMON
        # =========================
        sermon = (

            db.query(Sermon)

            .filter(
                Sermon.id == sermon_id
            )

            .filter(
                Sermon.author_id
                == current_user.id
            )

            .first()
        )

        if not sermon:

            raise HTTPException(
                status_code=404,
                detail="Sermon not found"
            )

        # =========================
        # VALIDATE RECIPIENT
        # =========================
        recipient = (

            db.query(User)

            .filter(
                User.id == to_pastor_id
            )

            .first()
        )

        if not recipient:

            raise HTTPException(
                status_code=404,
                detail="Recipient not found"
            )

        # =========================
        # MUST BE PASTOR
        # =========================
        if recipient.role != "pastor":

            raise HTTPException(
                status_code=400,
                detail="Recipient must be a pastor"
            )

        # =========================
        # PREVENT SELF SHARE
        # =========================
        if recipient.id == current_user.id:

            raise HTTPException(
                status_code=400,
                detail="Cannot share with yourself"
            )

        # =========================
        # PREVENT DUPLICATES
        # =========================
        existing = (

            db.query(SharedSermon)

            .filter(
                SharedSermon.sermon_id
                == sermon.id
            )

            .filter(
                SharedSermon.to_pastor_id
                == recipient.id
            )

            .first()
        )

        if existing:

            return {

                "success": True,

                "message":
                    "Sermon already shared"
            }

        # =========================
        # CREATE SHARE
        # =========================
        shared = SharedSermon(

            sermon_id=sermon.id,

            from_pastor_id=
                current_user.id,

            to_pastor_id=
                recipient.id,

            comment=comment
        )

        # =========================
        # ANALYTICS
        # =========================
        sermon.shares = (
            sermon.shares or 0
        ) + 1

        # =========================
        # SAVE
        # =========================
        db.add(shared)

        db.commit()

        return {

            "success": True,

            "message":
                "Sermon shared successfully"
        }

    except HTTPException:

        raise

    except Exception as e:

        print(
            "SHARE SERMON ERROR:",
            str(e)
        )

        raise HTTPException(
            status_code=500,
            detail=
                "Failed to share sermon"
        )
# =========================================
# SEARCH PASTORS
# =========================================
@router.get("/pastors/search")
async def search_pastors(

    query: str,

    db: Session = Depends(
        get_db
    ),

    current_user: User = Depends(
        get_current_user
    )
):

    try:

        # =========================
        # SEARCH USERS
        # =========================
        pastors = (

            db.query(User)

            .filter(
                User.role == "pastor"
            )

            .filter(
                User.name.ilike(
                    f"%{query}%"
                )
            )

            .limit(10)

            .all()
        )

        # =========================
        # FORMAT
        # =========================
        results = []

        for pastor in pastors:

            results.append({

                "id":
                    pastor.id,

                "name":
                    pastor.name,

                "email":
                    pastor.email
            })

        return {

            "success": True,

            "pastors": results
        }

    except Exception as e:

        print(
            "SEARCH PASTORS ERROR:",
            str(e)
        )

        raise HTTPException(
            status_code=500,
            detail=
              "Failed to search pastors"
        )
    
# =========================================
# SHARED SERMONS INBOX
# =========================================
@router.get("/shared-sermons")
async def shared_with_me(

    db: Session = Depends(
        get_db
    ),

    current_user: User = Depends(
        get_current_user
    )
):

    try:

        # =========================
        # PASTOR ONLY
        # =========================
        if current_user.role not in [
            "pastor",
            "admin"
        ]:

            raise HTTPException(
                status_code=403,
                detail="Pastor verification required"
            )

        # =========================
        # FETCH SHARED SERMONS
        # =========================
        shared = (

            db.query(SharedSermon)

            .filter(
                SharedSermon.to_pastor_id
                == current_user.id
            )

            .order_by(
                SharedSermon.created_at.desc()
            )

            .all()
        )

        # =========================
        # FORMAT RESPONSE
        # =========================
        results = []

        for item in shared:

            sermon = item.sermon

            sender = item.from_pastor

            results.append({

                "share_id":
                    item.id,

                "sermon_id":
                    sermon.id,

                "title":
                    sermon.title,

                "scripture":
                    sermon.scripture,

                "comment":
                    item.comment,

                "shared_at":
                    item.created_at,

                "sender_name":
                    sender.name
                    if sender
                    else "Unknown",

                "sender_email":
                    sender.email
                    if sender
                    else ""
            })

        return {

            "success": True,

            "shared_sermons":
                results
        }

    except HTTPException:

        raise

    except Exception as e:

        print(
            "SHARED INBOX ERROR:",
            str(e)
        )

        raise HTTPException(
            status_code=500,
            detail=
                "Failed to load shared sermons"
        )
    
# =========================================
# ADD SERMON COMMENT
# =========================================
@router.post("/sermon/comment")
async def add_sermon_comment(

    payload: dict,

    db: Session = Depends(
        get_db
    ),

    current_user: User = Depends(
        get_current_user
    )
):

    try:

        # =========================
        # PASTOR ONLY
        # =========================
        if current_user.role not in [
            "pastor",
            "admin"
        ]:

            raise HTTPException(
                status_code=403,
                detail="Pastor verification required"
            )

        # =========================
        # INPUTS
        # =========================
        sermon_id = payload.get(
            "sermon_id"
        )

        comment_text = payload.get(
            "comment"
        )

        # =========================
        # VALIDATE INPUTS
        # =========================
        if not sermon_id:

            raise HTTPException(
                status_code=400,
                detail="sermon_id is required"
            )

        if (
            not comment_text
            or not comment_text.strip()
        ):

            raise HTTPException(
                status_code=400,
                detail="Comment required"
            )

        # =========================
        # VALIDATE SERMON
        # =========================
        sermon = (

            db.query(Sermon)

            .filter(
                Sermon.id == sermon_id
            )

            .first()
        )

        if not sermon:

            raise HTTPException(
                status_code=404,
                detail="Sermon not found"
            )

        # =========================
        # CREATE COMMENT
        # =========================
        comment = SermonComment(

            sermon_id=sermon.id,

            pastor_id=current_user.id,

            comment=comment_text.strip()
        )

        db.add(comment)

        db.commit()

        db.refresh(comment)

        return {

            "success": True,

            "message":
                "Comment added successfully",

            "comment_id":
                comment.id
        }

    except HTTPException:

        raise

    except Exception as e:

        print(
            "ADD COMMENT ERROR:",
            str(e)
        )

        raise HTTPException(
            status_code=500,
            detail=
                "Failed to add comment"
        )
    
# =========================================
# GET SERMON COMMENTS
# =========================================
@router.get("/sermon/comments/{sermon_id}")
async def get_sermon_comments(

    sermon_id: int,

    db: Session = Depends(
        get_db
    ),

    current_user: User = Depends(
        get_current_user
    )
):

    try:

        # =========================
        # PASTOR ONLY
        # =========================
        if current_user.role not in [
            "pastor",
            "admin"
        ]:

            raise HTTPException(
                status_code=403,
                detail="Pastor verification required"
            )

        # =========================
        # VALIDATE SERMON
        # =========================
        sermon = (

            db.query(Sermon)

            .filter(
                Sermon.id == sermon_id
            )

            .first()
        )

        if not sermon:

            raise HTTPException(
                status_code=404,
                detail="Sermon not found"
            )

        # =========================
        # FETCH COMMENTS
        # =========================
        comments = (

            db.query(SermonComment)

            .filter(
                SermonComment.sermon_id
                == sermon_id
            )

            .order_by(
                SermonComment.created_at.asc()
            )

            .all()
        )

        # =========================
        # FORMAT RESPONSE
        # =========================
        results = []

        for item in comments:

            results.append({

                "id":
                    item.id,

                "comment":
                    item.comment,

                "created_at":
                    item.created_at,

                "pastor_name":
                    item.pastor.name
                    if item.pastor
                    else "Unknown"
            })

        return {

            "success": True,

            "comments":
                results
        }

    except HTTPException:

        raise

    except Exception as e:

        print(
            "GET COMMENTS ERROR:",
            str(e)
        )

        raise HTTPException(
            status_code=500,
            detail=
                "Failed to load comments"
        )